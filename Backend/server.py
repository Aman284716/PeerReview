from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Form, Header
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os
import uuid
import shutil
import subprocess
import json
import logging
from urllib.parse import urlparse
import numpy as np
from openai import AzureOpenAI
from datetime import datetime
import time
from docx import Document
import psutil
from dotenv import load_dotenv
import asyncio
from concurrent.futures import ThreadPoolExecutor
from functools import partial
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)
load_dotenv()

app = FastAPI(
    title="Code Review System - AI Standards Scanner",
    description="API to scan a public GitHub repository against industry and company-specific standards for code and IaC files using Azure OpenAI and RAG, with standards from .docx files.",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Vite dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
UPLOAD_FOLDER = Path("uploads")
INDUSTRY_STANDARDS_DIR = Path(r"C:\Users\286194\Documents\Aman\code_review\PeerReview\Backend\standards\industry_standards")
COMPANY_STANDARDS_DIR = Path(r"C:\Users\286194\Documents\Aman\code_review\PeerReview\Backend\standards\company_specific")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
INDUSTRY_STANDARDS_DIR.mkdir(exist_ok=True)
COMPANY_STANDARDS_DIR.mkdir(exist_ok=True)
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_MODEL = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
AZURE_OPENAI_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
MAX_CONCURRENT_API_CALLS = 10
MAX_WORKERS = 20

# Initialize Azure OpenAI client
try:
    openai_client = AzureOpenAI(
        api_key=AZURE_OPENAI_KEY,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_version=AZURE_OPENAI_VERSION
    )
    logger.debug("Azure OpenAI client initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize Azure OpenAI client: {e}")
    openai_client = None

security = HTTPBearer()

def read_docx_content(file_path: Path):
    """Read text content from a .docx file."""
    try:
        doc = Document(file_path)
        return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        logger.error(f"Error reading {file_path.name}: {str(e)}")
        return f"Error reading {file_path.name}: {str(e)}"

def analyze_doc_content(content: str):
    """Analyze .docx content using Azure OpenAI."""
    if not openai_client:
        return "Azure OpenAI client not initialized."
    try:
        prompt = f"Summarize the key rules or standards in the following document content and identify any compliance issues:\n\n{content[:4000]}"
        response = openai_client.chat.completions.create(
            model=AZURE_OPENAI_MODEL or "gpt-35-turbo",
            messages=[
                {"role": "system", "content": "You are an expert in analyzing coding standards and compliance."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=500,
            temperature=0.5
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"Azure OpenAI analysis failed: {str(e)}")
        return f"Analysis failed: {str(e)}"

def read_docx(file_path):
    """Read text from a .docx file (used for GitHub scanning)."""
    logger.debug(f"Reading .docx file: {file_path}")
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if not content:
            logger.warning(f"File {file_path} is empty")
            return None
        logger.debug(f"Extracted {len(content)} characters from {file_path}")
        return content
    except Exception as e:
        logger.error(f"Failed to read .docx file {file_path}: {e}")
        return None

# Load standards from .docx files
INDUSTRY_STANDARDS_FILE = INDUSTRY_STANDARDS_DIR / "Industry_Standards-updated.docx"
COMPANY_STANDARDS_FILE = COMPANY_STANDARDS_DIR / "company_specific_standards.docx"
INDUSTRY_STANDARDS_DOC = read_docx(INDUSTRY_STANDARDS_FILE)
COMPANY_STANDARDS_DOC = read_docx(COMPANY_STANDARDS_FILE)

if not INDUSTRY_STANDARDS_DOC:
    logger.warning("Industry standards document not found or empty, using fallback")
    INDUSTRY_STANDARDS_DOC = """
    1. OWASP Top 10: Ensure no SQL injection vulnerabilities (A03:2021).
    2. CWE Top 25: Avoid improper input validation (CWE-20).
    3. Secure Coding: Use HTTPS for network communications.
    4. Terraform: Use least privilege for IAM roles and validate inputs.
    5. YAML/CloudFormation: Ensure secure S3 bucket configurations.
    6. Bash: Avoid unquoted variables to prevent command injection.
    7. Dockerfile: Use minimal base images and avoid running as root.
    """

if not COMPANY_STANDARDS_DOC:
    logger.warning("Company standards document not found or empty, using fallback")
    COMPANY_STANDARDS_DOC = """
    1. Code Style: Use snake_case for variable names in Python.
    2. Security: Avoid hardcoding sensitive data like API keys.
    3. Documentation: Include docstrings for all functions.
    4. Terraform: Use modules for reusable configurations.
    5. YAML: Maintain consistent indentation and validate schemas.
    6. Bash: Include shebang and error handling.
    7. Dockerfile: Pin image versions and include MAINTAINER labels.
    """

logger.info(f"Industry standards loaded: {len(INDUSTRY_STANDARDS_DOC)} characters")
logger.info(f"Company standards loaded: {len(COMPANY_STANDARDS_DOC)} characters")

async def verify_admin_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    token = credentials.credentials
    if token != "token_for_demo":
        raise HTTPException(status_code=401, detail="Invalid or missing admin token")
    return token

def read_docx_rules(file_path: Path):
    try:
        doc = Document(file_path)
        rules = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return rules
    except Exception as e:
        return [f"Error reading {file_path.name}: {str(e)}"]

def generate_unique_id():
    return str(uuid.uuid4())

def safe_rmtree(path, retries=5, delay=2):
    logger.debug(f"Attempting to remove directory: {path}")
    for attempt in range(retries):
        try:
            for proc in psutil.process_iter(['name']):
                if proc.info['name'].lower().startswith('git'):
                    proc.terminate()
                    proc.wait(timeout=3)
                    logger.debug(f"Terminated Git process: {proc.pid}")
            shutil.rmtree(path, ignore_errors=False)
            logger.debug(f"Directory {path} removed successfully")
            return True
        except PermissionError as e:
            logger.warning(f"PermissionError on attempt {attempt + 1}: {e}")
            time.sleep(delay)
        except Exception as e:
            logger.error(f"Error removing directory {path}: {e}")
            time.sleep(delay)
    logger.error(f"Failed to remove directory {path} after {retries} attempts")
    return False

def clone_repo(github_url):
    repo_dir = os.path.join(UPLOAD_FOLDER, f"repo_{generate_unique_id()}")
    repo_dir = repo_dir.replace("/", os.sep)
    logger.debug(f"Cloning repository from {github_url} to {repo_dir}")
    try:
        result = subprocess.run(
            ["git", "clone", github_url, repo_dir],
            capture_output=True, text=True, timeout=120, check=True
        )
        logger.debug(f"Clone successful: {result.stdout}")
        return repo_dir
    except subprocess.TimeoutExpired as e:
        logger.error(f"Clone timed out: {e}")
        return None
    except subprocess.CalledProcessError as e:
        logger.error(f"Failed to clone repository: {e.stderr}")
        return None
    except FileNotFoundError:
        logger.error("Git is not installed or not found in PATH")
        return None

def read_code_files(repo_dir, max_files=1000, max_size=10000):
    logger.debug(f"Reading code and IaC files from {repo_dir}")
    code_files = []
    extensions = (
        ".py", ".js", ".java", ".cpp", ".c", ".go", ".html", ".css", ".cs",
        ".php", ".rb", ".kt", ".swift", ".scala", ".ts", ".hcl", ".yaml", ".yml", ".sh"
    )
    dockerfile_names = ("Dockerfile",)
    for root, _, files in os.walk(repo_dir):
        for file in files:
            if len(code_files) >= max_files:
                logger.debug(f"Reached max_files limit: {max_files}")
                break
            file_path = os.path.join(root, file)
            if file.endswith(extensions) or file in dockerfile_names:
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read(max_size)
                        code_files.append((file_path, content))
                        logger.debug(f"Read file: {file_path}")
                except Exception as e:
                    logger.warning(f"Failed to read {file_path}: {e}")
    if not code_files:
        logger.warning(f"No code or IaC files found in {repo_dir}")
    else:
        logger.debug(f"Found {len(code_files)} code and IaC files")
    return code_files

def simple_vectorize(text):
    if not text or not text.strip():
        return np.array([0.0])
    words = text.lower().split()
    word_count = len(words)
    char_count = len(text)
    unique_words = len(set(words))
    return np.array([word_count, char_count, unique_words], dtype=float)

def retrieve_relevant_sections(doc_content, query, top_k=3):
    logger.debug("Retrieving relevant sections for RAG")
    try:
        if not doc_content or not doc_content.strip():
            logger.warning("Document content is empty")
            return []
        sections = []
        paragraphs = [p.strip() for p in doc_content.split('\n\n') if p.strip()]
        for paragraph in paragraphs:
            lines = [line.strip() for line in paragraph.split('\n') if line.strip()]
            sections.extend(lines)
        sections = [s for s in sections if len(s) > 20]
        if not sections:
            logger.warning("No valid sections found in document")
            return []
        query_vector = simple_vectorize(query)
        section_vectors = [simple_vectorize(section) for section in sections]
        similarities = []
        for sv in section_vectors:
            if np.linalg.norm(query_vector) == 0 or np.linalg.norm(sv) == 0:
                similarities.append(0.0)
            else:
                sim = np.dot(query_vector, sv) / (np.linalg.norm(query_vector) * np.linalg.norm(sv))
                similarities.append(sim)
        if len(similarities) == 0:
            return []
        top_indices = np.argsort(similarities)[-top_k:]
        relevant_sections = [sections[i] for i in top_indices if similarities[i] > 0.1]
        logger.debug(f"Retrieved {len(relevant_sections)} relevant sections")
        return relevant_sections
    except Exception as e:
        logger.error(f"RAG retrieval failed: {e}")
        return []

def get_language_specific_standards(file_path, standards_doc):
    file_name = os.path.basename(file_path)
    extension = os.path.splitext(file_name)[1].lower()
    language_map = {
        '.java': 'Java',
        '.js': 'JavaScript',
        '.ts': 'TypeScript',
        '.py': 'Python',
        '.cs': 'C#',
        '.cpp': 'C++',
        '.c': 'C',
        '.php': 'PHP',
        '.go': 'Go',
        '.rb': 'Ruby',
        '.kt': 'Kotlin',
        '.swift': 'Swift',
        '.scala': 'Scala',
        '.hcl': 'Terraform',
        '.yaml': 'YAML',
        '.yml': 'YAML',
        '.sh': 'Bash'
    }
    if file_name == 'Dockerfile':
        language = 'Dockerfile'
    else:
        language = language_map.get(extension, 'General')
    query = f"{language} coding standards best practices security"
    relevant_sections = retrieve_relevant_sections(standards_doc, query, top_k=5)
    return relevant_sections, language

async def analyze_file_industry(file_path, content, language, relevant_sections, standard_desc, semaphore):
    async with semaphore:
        try:
            system_message = f"""You are an expert code reviewer specializing in {language} development and IaC configurations. 
Analyze code or configuration files against {standard_desc} and provide a detailed analysis in JSON format.
Focus on security vulnerabilities, code quality issues, and adherence to industry standards."""
            user_message = f"""
Analyze the following {language} code or configuration against {standard_desc}:

File: {os.path.basename(file_path)}

Content:
```{language.lower()}
{content[:2000]}
```

Relevant Industry Standards:
{chr(10).join(relevant_sections[:3])}

Provide a comprehensive analysis covering:
1. Security vulnerabilities (OWASP, CWE)
2. Code quality issues (or configuration best practices)
3. Industry standard violations
4. Recommendations for improvement

Return your response in this exact JSON format:
{{
  "file": "{os.path.basename(file_path)}",
  "language": "{language}",
  "security_issues": "List any security vulnerabilities found",
  "quality_issues": "List code or configuration quality problems",
  "standard_violations": "List industry standard violations",
  "recommendations": "Provide specific improvement recommendations",
  "severity": "HIGH/MEDIUM/LOW based on issues found",
  "standards_applied": {json.dumps(relevant_sections[:3])}
}}
"""
            logger.debug(f"Sending chat completion request for {file_path}")
            response = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: openai_client.chat.completions.create(
                    model=AZURE_OPENAI_MODEL,
                    messages=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message}
                    ],
                    max_tokens=800,
                    temperature=0.2,
                    response_format={"type": "json_object"}
                )
            )
            response_text = response.choices[0].message.content.strip()
            logger.debug(f"Raw AI response for {file_path}: {response_text[:200]}...")
            try:
                analysis = json.loads(response_text)
                analysis["file"] = file_path
                analysis["full_path"] = file_path
                logger.debug(f"Successfully processed {file_path}")
                return analysis
            except json.JSONDecodeError as json_err:
                logger.warning(f"Failed to parse AI response for {file_path}: {json_err}")
                return {
                    "file": file_path,
                    "language": language,
                    "issues": response_text,
                    "error": "Failed to parse AI response",
                    "standards_applied": relevant_sections
                }
        except Exception as e:
            logger.error(f"Azure OpenAI call failed for {file_path}: {e}")
            return {
                "file": file_path,
                "issues": f"Analysis failed: {str(e)}",
                "error": str(e),
                "standards_applied": []
            }

async def industry_standards_agent(repo_dir, industrial_standard):
    logger.debug(f"Starting industry standards scan with standard: {industrial_standard}")
    if not openai_client:
        logger.error("Azure OpenAI client not initialized")
        return None
    try:
        code_files = read_code_files(repo_dir)
        if not code_files:
            logger.warning("No code or IaC files found for industry standards scan")
            return {"error": "No code or IaC files found"}
        standards_map = {
            "owasp": "OWASP Top 10 security standards",
            "cwe": "CWE Top 25 most dangerous software weaknesses",
            "sonarqube": "SonarQube code quality standards",
            "general": "General industry coding standards"
        }
        standard_desc = standards_map.get(industrial_standard, "General industry coding standards")
        logger.debug(f"Using standard description: {standard_desc}")
        semaphore = asyncio.Semaphore(MAX_CONCURRENT_API_CALLS)
        tasks = []
        for file_path, content in code_files:
            relevant_sections, language = get_language_specific_standards(file_path, INDUSTRY_STANDARDS_DOC)
            if not relevant_sections:
                query = f"Check this {language} code for {standard_desc} security issues quality problems"
                relevant_sections = retrieve_relevant_sections(INDUSTRY_STANDARDS_DOC, query, top_k=3)
            if not relevant_sections:
                logger.warning(f"No relevant industry standards found for {file_path}")
                relevant_sections = ["General coding best practices apply"]
            tasks.append(
                analyze_file_industry(file_path, content, language, relevant_sections, standard_desc, semaphore)
            )
        findings = await asyncio.gather(*tasks, return_exceptions=True)
        findings = [f for f in findings if not isinstance(f, Exception)]
        return {
            "findings": findings,
            "standard": standard_desc,
            "total_files_analyzed": len(findings),
            "summary": f"Analyzed {len(findings)} files against {standard_desc}"
        }
    except Exception as e:
        logger.error(f"Industry standards scan failed: {e}")
        return {"error": f"Industry standards scan failed: {str(e)}"}

async def analyze_file_company(file_path, content, language, relevant_sections, semaphore):
    async with semaphore:
        try:
            system_message = f"""You are an expert code reviewer focusing on company-specific coding standards and style guidelines for {language}.
Analyze code or configuration files for adherence to company policies, coding conventions, and internal best practices."""
            user_message = f"""
Analyze the following {language} code or configuration against company-specific standards:

File: {os.path.basename(file_path)}

Content:
```{language.lower()}
{content[:2000]}
```

Relevant Company Standards:
{chr(10).join(relevant_sections[:3])}

Focus on:
1. Coding style and conventions
2. Documentation standards
3. Company-specific security policies
4. Internal best practices

Return your response in this exact JSON format:
{{
  "file": "{os.path.basename(file_path)}",
  "language": "{language}",
  "style_issues": "List coding style violations",
  "documentation_issues": "List documentation problems",
  "policy_violations": "List company policy violations",
  "recommendations": "Provide specific improvement recommendations",
  "compliance_score": "Score from 1-10 for company standard compliance",
  "standards_applied": {json.dumps(relevant_sections[:3])}
}}
"""
            logger.debug(f"Sending chat completion request for {file_path}")
            response = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: openai_client.chat.completions.create(
                    model=AZURE_OPENAI_MODEL,
                    messages=[
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message}
                    ],
                    max_tokens=800,
                    temperature=0.2,
                    response_format={"type": "json_object"}
                )
            )
            response_text = response.choices[0].message.content.strip()
            logger.debug(f"Raw AI response for {file_path}: {response_text[:200]}...")
            try:
                analysis = json.loads(response_text)
                analysis["file"] = file_path
                analysis["full_path"] = file_path
                logger.debug(f"Successfully processed {file_path}")
                return analysis
            except json.JSONDecodeError as json_err:
                logger.warning(f"Failed to parse AI response for {file_path}: {json_err}")
                return {
                    "file": file_path,
                    "language": language,
                    "issues": response_text,
                    "error": "Failed to parse AI response",
                    "standards_applied": relevant_sections
                }
        except Exception as e:
            logger.error(f"Azure OpenAI call failed for {file_path}: {e}")
            return {
                "file": file_path,
                "issues": f"Analysis failed: {str(e)}",
                "error": str(e),
                "standards_applied": []
            }

async def company_specific_agent(repo_dir):
    logger.debug("Starting company-specific standards scan")
    if not openai_client:
        logger.error("Azure OpenAI client not initialized")
        return None
    try:
        code_files = read_code_files(repo_dir)
        if not code_files:
            logger.warning("No code or IaC files found for company-specific scan")
            return {"error": "No code or IaC files found"}
        semaphore = asyncio.Semaphore(MAX_CONCURRENT_API_CALLS)
        tasks = []
        for file_path, content in code_files:
            relevant_sections, language = get_language_specific_standards(file_path, COMPANY_STANDARDS_DOC)
            if not relevant_sections:
                query = f"Check this {language} code or configuration for company-specific coding standards style guidelines"
                relevant_sections = retrieve_relevant_sections(COMPANY_STANDARDS_DOC, query, top_k=3)
            if not relevant_sections:
                logger.warning(f"No relevant company standards found for {file_path}")
                relevant_sections = ["General company coding guidelines apply"]
            tasks.append(
                analyze_file_company(file_path, content, language, relevant_sections, semaphore)
            )
        findings = await asyncio.gather(*tasks, return_exceptions=True)
        findings = [f for f in findings if not isinstance(f, Exception)]
        return {
            "findings": findings,
            "standard": "Company-specific standards",
            "total_files_analyzed": len(findings),
            "summary": f"Analyzed {len(findings)} files against company standards"
        }
    except Exception as e:
        logger.error(f"Company-specific scan failed: {e}")
        return {"error": f"Company-specific scan failed: {str(e)}"}

def generate_report(industry_results, company_results):
    logger.debug("Generating comprehensive report")
    industry_files = len(industry_results.get("findings", [])) if industry_results else 0
    company_files = len(company_results.get("findings", [])) if company_results else 0
    high_severity_issues = 0
    medium_severity_issues = 0
    low_severity_issues = 0
    if industry_results and "findings" in industry_results:
        for finding in industry_results["findings"]:
            severity = finding.get("severity", "MEDIUM").upper()
            if severity == "HIGH":
                high_severity_issues += 1
            elif severity == "MEDIUM":
                medium_severity_issues += 1
            else:
                low_severity_issues += 1
    report = {
        "scan_summary": {
            "timestamp": datetime.now().isoformat(),
            "industry_files_analyzed": industry_files,
            "company_files_analyzed": company_files,
            "total_files_analyzed": max(industry_files, company_files),
            "high_severity_issues": high_severity_issues,
            "medium_severity_issues": medium_severity_issues,
            "low_severity_issues": low_severity_issues
        },
        "industry_standards_analysis": industry_results or {"error": "Industry standards scan failed"},
        "company_standards_analysis": company_results or {"error": "Company-specific scan failed"},
        "recommendations": {
            "priority_actions": [],
            "general_improvements": []
        }
    }
    if high_severity_issues > 0:
        report["recommendations"]["priority_actions"].append(f"Address {high_severity_issues} high-severity security/quality issues immediately")
    if industry_results and company_results:
        report["recommendations"]["general_improvements"].append("Review and implement coding and IaC standards consistently across all files")
        report["recommendations"]["general_improvements"].append("Consider implementing automated code and IaC quality checks in CI/CD pipeline")
    return report

async def process_github_repository(github_url, industrial_standard):
    logger.info(f"Processing GitHub repository: {github_url}")
    parsed_url = urlparse(github_url)
    if not parsed_url.netloc.endswith("github.com"):
        logger.error("Invalid GitHub URL provided")
        raise ValueError("Invalid GitHub URL")
    if not INDUSTRY_STANDARDS_DOC:
        logger.error("Industry standards document is empty or failed to load")
        raise RuntimeError("Failed to load industry standards")
    if not COMPANY_STANDARDS_DOC:
        logger.error("Company standards document is empty or failed to load")
        raise RuntimeError("Failed to load company standards")
    if not openai_client:
        logger.error("Azure OpenAI client is not initialized")
        raise RuntimeError("Azure OpenAI client not initialized")
    repo_dir = clone_repo(github_url)
    if not repo_dir:
        logger.error("Repository cloning failed")
        raise RuntimeError("Failed to clone repository")
    try:
        logger.info("Running industry and company-specific agents")
        industry_results = await industry_standards_agent(repo_dir, industrial_standard)
        company_results = await company_specific_agent(repo_dir)
        report = generate_report(industry_results, company_results)
        output_file = os.path.join(UPLOAD_FOLDER, f"report_{generate_unique_id()}.json")
        output_file = output_file.replace("/", os.sep)
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        logger.info(f"Report saved to {output_file}")
        return output_file
    except Exception as e:
        logger.error(f"Processing failed: {e}")
        raise
    finally:
        if os.path.exists(repo_dir):
            safe_rmtree(repo_dir)

class StandardsScanInput(BaseModel):
    github_url: str
    industrial_standard: str

    class Config:
        schema_extra = {
            "example": {
                "github_url": "https://github.com/pallets/flask.git",
                "industrial_standard": "owasp"
            }
        }

class StandardsScanResponse(BaseModel):
    message: str
    download_url: str

    class Config:
        schema_extra = {
            "example": {
                "message": "Repository scanned successfully",
                "download_url": "/download/report_123e4567-e89b-12d3-a456-426614174000.json"
            }
        }

@app.get("/")
async def root():
    return {
        "message": "Code Review System - AI Standards Scanner",
        "version": "1.0.0",
        "endpoints": {
            "/scan-standards": "POST - Scan repository against standards",
            "/download/{filename}": "GET - Download generated report",
            "/standards": "GET - List standards",
            "/upload-standard": "POST - Upload standard document",
            "/scan-standards-documents": "POST - Scan standards documents",
            "/docs": "GET - API documentation"
        }
    }

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "openai_client": "initialized" if openai_client else "not initialized",
        "industry_standards": "loaded" if INDUSTRY_STANDARDS_DOC else "not loaded",
        "company_standards": "loaded" if COMPANY_STANDARDS_DOC else "not loaded"
    }

@app.post("/scan-standards", response_model=StandardsScanResponse)
async def scan_standards(scan_input: StandardsScanInput):
    logger.info(f"Received scan request: {scan_input.dict()}")
    try:
        output_file = await process_github_repository(
            scan_input.github_url,
            scan_input.industrial_standard
        )
        logger.info(f"Scan completed successfully: {output_file}")
        return StandardsScanResponse(
            message="Repository scanned successfully",
            download_url=f"/download/{os.path.basename(output_file)}"
        )
    except ValueError as e:
        logger.error(f"Validation error: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Scan error: {e}")
        raise HTTPException(status_code=500, detail=f"Standards scanning error: {str(e)}")

# Replace /standards in server.py
@app.get("/standards")
async def get_standards(token: str = Depends(verify_admin_token)):
    language_rules = {
        "General": {"industry": [], "company": []},
        "Python": {"industry": [], "company": []},
        "JavaScript": {"industry": [], "company": []},
        "Java": {"industry": [], "company": []},
        "C#": {"industry": [], "company": []},
        "Terraform": {"industry": [], "company": []},
        "YAML": {"industry": [], "company": []},
        "Bash": {"industry": [], "company": []},
        "Dockerfile": {"industry": [], "company": []}
    }
    for file_path in INDUSTRY_STANDARDS_DIR.glob("*.docx"):
        rules = read_docx_rules(file_path)
        for rule in rules:
            rule_data = {"text": rule, "file": file_path.name}
            rule_lower = rule.lower()
            if "python" in rule_lower:
                language_rules["Python"]["industry"].append(rule_data)
            elif "javascript" in rule_lower or "js" in rule_lower or "node" in rule_lower:
                language_rules["JavaScript"]["industry"].append(rule_data)
            elif "java" in rule_lower:
                language_rules["Java"]["industry"].append(rule_data)
            elif "c#" in rule_lower or "csharp" in rule_lower:
                language_rules["C#"]["industry"].append(rule_data)
            elif "terraform" in rule_lower:
                language_rules["Terraform"]["industry"].append(rule_data)
            elif "yaml" in rule_lower or "cloudformation" in rule_lower:
                language_rules["YAML"]["industry"].append(rule_data)
            elif "bash" in rule_lower or "shell" in rule_lower:
                language_rules["Bash"]["industry"].append(rule_data)
            elif "dockerfile" in rule_lower or "docker" in rule_lower:
                language_rules["Dockerfile"]["industry"].append(rule_data)
            else:
                language_rules["General"]["industry"].append(rule_data)
    for file_path in COMPANY_STANDARDS_DIR.glob("*.docx"):
        rules = read_docx_rules(file_path)
        for rule in rules:
            rule_data = {"text": rule, "file": file_path.name}
            rule_lower = rule.lower()
            if "python" in rule_lower:
                language_rules["Python"]["company"].append(rule_data)
            elif "javascript" in rule_lower or "js" in rule_lower or "node" in rule_lower:
                language_rules["JavaScript"]["company"].append(rule_data)
            elif "java" in rule_lower:
                language_rules["Java"]["company"].append(rule_data)
            elif "c#" in rule_lower or "csharp" in rule_lower:
                language_rules["C#"]["company"].append(rule_data)
            elif "terraform" in rule_lower:
                language_rules["Terraform"]["company"].append(rule_data)
            elif "yaml" in rule_lower:
                language_rules["YAML"]["company"].append(rule_data)
            elif "bash" in rule_lower or "shell" in rule_lower:
                language_rules["Bash"]["company"].append(rule_data)
            elif "dockerfile" in rule_lower or "docker" in rule_lower:
                language_rules["Dockerfile"]["company"].append(rule_data)
            else:
                language_rules["General"]["company"].append(rule_data)
    return language_rules

@app.post("/upload-standard")
async def upload_standard(
    file: UploadFile = File(...),
    standard_type: str = Form(...),
    token: str = Depends(verify_admin_token)
):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    target_dir = INDUSTRY_STANDARDS_DIR if standard_type == "industry" else COMPANY_STANDARDS_DIR
    file_path = target_dir / file.filename
    try:
        with file_path.open("wb") as f:
            shutil.copyfileobj(file.file, f)
        return {"message": "File uploaded successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")

@app.post("/scan-standards-documents")
async def scan_standards_documents(token: str = Depends(verify_admin_token)):
    try:
        industry_files = list(INDUSTRY_STANDARDS_DIR.glob("*.docx"))
        company_files = list(COMPANY_STANDARDS_DIR.glob("*.docx"))
        industry_count = len(industry_files)
        company_count = len(company_files)
        total_standards = industry_count + company_count
        analysis_results = []
        for file_path in industry_files + company_files:
            content = read_docx_content(file_path)
            if content.startswith("Error"):
                analysis_results.append({
                    "file": file_path.name,
                    "type": "industry" if file_path in industry_files else "company",
                    "status": "error",
                    "details": content
                })
            else:
                analysis = analyze_doc_content(content)
                analysis_results.append({
                    "file": file_path.name,
                    "type": "industry" if file_path in industry_files else "company",
                    "status": "success",
                    "details": analysis
                })
        return {
            "total_standards": total_standards,
            "industry_standards_count": industry_count,
            "company_standards_count": company_count,
            "details": "Scanned all standards documents successfully.",
            "analysis_results": analysis_results
        }
    except Exception as e:
        logger.error(f"Scan failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Scan failed: {str(e)}")

@app.get("/download/{filename}")
async def download_file(filename: str):
    logger.debug(f"Download request for {filename}")
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if os.path.exists(file_path):
        logger.debug(f"Serving file: {file_path}")
        return FileResponse(
            file_path,
            media_type="application/json",
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    logger.error(f"File not found: {file_path}")
    raise HTTPException(status_code=404, detail="File not found")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)